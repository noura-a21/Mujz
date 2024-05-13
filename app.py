import os
from flask import Flask, render_template, request, session, redirect, url_for, flash, jsonify
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
import re
from flask import render_template
from docx import Document
from PyPDF2 import PdfReader
import requests
import json
from pydub import AudioSegment
import azure.cognitiveservices.speech as speechsdk
from langdetect import detect, LangDetectException

speech_key = "ca887f50f22c461c8b175b681faa791f"
service_region = "eastus"

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users1.sqlite3'
app.config['SECRET_KEY'] = 'your_secret_key'
db = SQLAlchemy(app)

first_request_processed = False
app.config['UPLOAD_FOLDER'] = 'upload'

ALLOWED_EXTENSIONS = {'doc', 'docx', 'pdf', 'mp3', 'wav'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

class User(db.Model):
   name = db.Column(db.String(100), nullable=False)
   email = db.Column(db.String(100), primary_key=True, nullable=False)
   password = db.Column(db.String(200), nullable=False)

# Initialize Database
def before_first_request():
    global first_request_processed
    if not first_request_processed:
        first_request_processed = True
        create_tables()
        db.create_all()

def create_tables():
    before_first_request()
    db.create_all()

global_var = ""
def modify_global(text):
    global global_var
    global_var = text

def is_english(text):
    try:
        # Detect the language of the text
        return detect(text) == 'en'
    except LangDetectException:
        return False
    
openai_api_key = "sk-proj-FM4bY9TeLD7jyD55lS3OT3BlbkFJCF5i1zqexAiWPKDCkgZ0"
# Function to summarize text using ChatGPT
def summarize_text_with_chatgpt(text):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai_api_key}"
    }
    data = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": f"Summarize this text: {text}"
            }
        ]
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        response_data = response.json()
        summary = response_data['choices'][0]['message']['content']
        return summary
    else:
        print("Error:", response.status_code, response.text)
        return None

def chat_with_chatgpt(text):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai_api_key}"
    }
    data = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": f"Read this text, analyze it well, and write to me a set of questions about this attached text, and I will discuss this text with you in the upcoming responses the text is: {text}"
            }
        ]
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        response_data = response.json()
        summary = response_data['choices'][0]['message']['content']
        return summary
    else:
        print("Error:", response.status_code, response.text)
        return None
    
def transcribe_audio_azure_continuous(audio_filename):
    # Configure speech recognition for English language
    speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)
    speech_config.speech_recognition_language = "en-US"

    audio_config = speechsdk.audio.AudioConfig(filename=audio_filename)

    # Create a recognizer
    speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)

    all_results = []
    done = False

    def recognized_callback(event):
        if event.result.reason == speechsdk.ResultReason.RecognizedSpeech:
            all_results.append(event.result.text)

    def canceled_callback(event):
        nonlocal done
        if event.result.reason == speechsdk.ResultReason.EndOfStream:
            done = True

    def stop_callback(event):
        nonlocal done
        done = True

    speech_recognizer.recognized.connect(recognized_callback)
    speech_recognizer.canceled.connect(canceled_callback)
    speech_recognizer.session_stopped.connect(stop_callback)

    speech_recognizer.start_continuous_recognition()
    while not done:
        pass

    speech_recognizer.stop_continuous_recognition()
    return ' '.join(all_results)

def transcribe_audio_from_mp3(mp3_file_path):
    wav_file_path = mp3_file_path[:-4] + '.wav'
    AudioSegment.from_mp3(mp3_file_path).export(wav_file_path, format="wav")
    return transcribe_audio_azure_continuous(wav_file_path)

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/signUp', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']

        # Validate email
        if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
            flash('Invalid email address.')
            return redirect(url_for('signup'))

        # Validate password
        password_pattern = r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d).{8,}$'
        if not re.match(password_pattern, password):
            flash('Password must be at least 8 characters long and include uppercase and lowercase letters, and numbers.')
            return redirect(url_for('signup'))

        # Check email exists
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            flash('Email already registered. Did you forget your password?')
            return redirect(url_for('signup'))

        hashed_password = generate_password_hash(password)
        new_user = User(name=name, email=email, password=hashed_password)
        db.session.add(new_user)
        db.session.commit()

        flash('Sign up successful.')
        return redirect(url_for('signin'))

    return render_template('signUp.html')

@app.route('/signIn', methods=['GET', 'POST'])
def signin():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        user = User.query.filter_by(email=email).first()
        if user and check_password_hash(user.password, password):
            session['email'] = user.email
            session['name'] = user.name
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password')

    return render_template('signIn.html')

@app.route('/recover-password', methods=['GET', 'POST'])
def recover_password():
    if request.method == 'POST':
        email = request.form['email']
        user = User.query.filter_by(email=email).first()
        if user:
            # Redirect to the reset password page with the user's email as a query parameter
            return redirect(url_for('reset_password', email=email))
        else:
            flash('No account associated with this email.')
    return render_template('recover-password.html')

@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password():
    if request.method == 'POST':
        email = request.args.get('email')
        new_password = request.form['new_password']

        # Validate password strength
        password_pattern = r'^(?=.*[A-Z])(?=.*[a-z])(?=.*\d).{8,}$'
        if not re.match(password_pattern, new_password):
            flash('Password must be at least 8 characters long and include uppercase and lowercase letters, and numbers.')
            return redirect(url_for('reset_password', email=email))

        user = User.query.filter_by(email=email).first()
        if user:
            # Update the user's password
            user.password = generate_password_hash(new_password)
            db.session.commit()
            flash('Your password has been updated.')
            return redirect(url_for('signin'))  # Assuming you have a login route
        else:
            flash('An error occurred. Please try the link from your email again.')
    else:
        email = request.args.get('email', '')
    return render_template('reset-password.html', email=email)

@app.route('/upload-file', methods=['POST'])
def upload_file():
    if 'email' not in session:
        return redirect(url_for('signin'))
    
    action = request.form.get('action')
    if action == 'write':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(url_for('dashboard'))

        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(url_for('dashboard'))

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            if '.' not in filename:
                flash('Invalid file format')
                return redirect(url_for('dashboard'))
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            flash('The file was successfully attached')

            # Determine file extension and process accordingly
            file_extension = filename.rsplit('.', 1)[1].lower()

            if file_extension in {'doc', 'docx'}:
                # Process Word document
                document = Document(save_path)
                doc_text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
                if not is_english(doc_text):
                    flash('You must enter an English file')
                    return redirect(url_for('dashboard'))
                summary = chat_with_chatgpt(doc_text)
                modify_global(summary)
                if summary:
                    flash('Questions have been generated.')
                    return render_template('dashboard.html', outputtext=summary)
                else:
                    flash('Failed .')

            elif file_extension == 'pdf':
                # Process PDF document
                pdf_reader = PdfReader(save_path)
                pdf_text = '\n'.join([page.extract_text() for page in pdf_reader.pages if page.extract_text() is not None])
                if not is_english(pdf_text):
                    flash('You must enter an English file')
                    return redirect(url_for('dashboard'))
                summary = chat_with_chatgpt(pdf_text)
                modify_global(summary)
                if summary:
                    flash('Questions have been generated.')
                    return render_template('dashboard.html', outputtext=summary)
                else:
                    flash('Failed .')

            elif file_extension == 'wav':
                audio_text = transcribe_audio_azure_continuous(save_path)
                if audio_text:
                    if not is_english(audio_text):
                        flash('You must enter an English audio file')
                        return redirect(url_for('dashboard'))
                    flash('Audio file content transcribed.')
                    summary = chat_with_chatgpt(audio_text)
                    modify_global(summary)
                    if summary:
                        flash('Questions have been generated.')
                        return render_template('dashboard.html', outputtext=summary)
                    else:
                        flash('Failed .')
                else:
                    flash('Failed to transcribe audio file content.')

            elif file_extension == 'mp3':
                audio_text = transcribe_audio_from_mp3(save_path)
                if audio_text:
                    if not is_english(audio_text):
                        flash('You must enter an English audio file')
                        return redirect(url_for('dashboard'))
                    flash('Audio file content transcribed.')
                    summary = chat_with_chatgpt(audio_text)
                    modify_global(summary)
                    if summary:
                        flash('Questions have been generated.')
                        return render_template('dashboard.html', outputtext=summary)
                    else:
                        flash('Failed .')
                else:
                    flash('Failed to transcribe audio file content.')

            flash(f'File {filename} has been uploaded and saved.')
            return redirect(url_for('dashboard'))
    
    elif action == 'upload':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(url_for('dashboard'))

        file = request.files['file']

        if file.filename == '':
            flash('No selected file')
            return redirect(url_for('dashboard'))

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            if '.' not in filename:
                flash('Invalid file format')
                return redirect(url_for('dashboard'))
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            flash('The file was successfully attached')

            # Determine file extension and process accordingly
            file_extension = filename.rsplit('.', 1)[1].lower()

            if file_extension in {'doc', 'docx'}:
                # Process Word document
                document = Document(save_path)
                doc_text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
                if not is_english(doc_text):
                    flash('You must enter an English file')
                    return redirect(url_for('dashboard'))
                summary = summarize_text_with_chatgpt(doc_text)
                modify_global(summary)
                if summary:
                    flash('Word document content summarized.')
                    return render_template('dashboard.html', outputtext=summary)
                else:
                    flash('Failed to summarize Word document content.')

            elif file_extension == 'pdf':
                # Process PDF document
                pdf_reader = PdfReader(save_path)
                pdf_text = '\n'.join([page.extract_text() for page in pdf_reader.pages if page.extract_text() is not None])
                if not is_english(pdf_text):
                    flash('You must enter an English file')
                    return redirect(url_for('dashboard'))
                summary = summarize_text_with_chatgpt(pdf_text)
                modify_global(summary)
                if summary:
                    flash('PDF document content summarized.')
                    return render_template('dashboard.html', outputtext=summary)
                else:
                    flash('Failed to summarize PDF document content.')

            elif file_extension == 'wav':
                audio_text = transcribe_audio_azure_continuous(save_path)
                if audio_text:
                    if not is_english(audio_text):
                        flash('You must enter an English audio file')
                        return redirect(url_for('dashboard'))
                    flash('Audio file content transcribed.')
                    print(audio_text)
                    summary = summarize_text_with_chatgpt(audio_text)
                    modify_global(summary)
                    if summary:
                        flash('Audio file content summarized.')
                        return render_template('dashboard.html', outputtext=summary)
                    else:
                        flash('Failed to summarize audio file content.')
                else:
                    flash('Failed to transcribe audio file content.')

            elif file_extension == 'mp3':
                audio_text = transcribe_audio_from_mp3(save_path)
                if audio_text:
                    if not is_english(audio_text):
                        flash('You must enter an English audio file')
                        return redirect(url_for('dashboard'))
                    flash('Audio file content transcribed.')
                    summary = summarize_text_with_chatgpt(audio_text)
                    modify_global(summary)
                    if summary:
                        flash('Audio file content summarized.')
                        return render_template('dashboard.html', outputtext=summary)
                    else:
                        flash('Failed to summarize audio file content.')
                else:
                    flash('Failed to transcribe audio file content.')

            flash(f'File {filename} has been uploaded and saved.')
            return redirect(url_for('dashboard'))

    else:
        flash('Invalid action.')
        return redirect(url_for('dashboard'))

@app.route('/questions', methods=['GET', 'POST'])
def questions():
    global global_var
    if request.method == 'GET':
        question = request.args.get('question')
        if question:
            full_question = global_var + " " + question
            chatgpt_response = chat_with_chatgpt_for_text(full_question)
            if chatgpt_response:
                return render_template('dashboard.html', outputtext=chatgpt_response)
            else:
                return render_template('dashboard.html', outputtext="Failed to get response from ChatGPT.")
        else:
            return render_template('dashboard.html', outputtext="No question provided.")
    else:
        return render_template('dashboard.html')

def chat_with_chatgpt_for_text(question):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {openai_api_key}"
    }
    data = {
        "model": "gpt-3.5-turbo",
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": question
            }
        ]
    }
    response = requests.post(url, headers=headers, json=data)
    if response.status_code == 200:
        response_data = response.json()
        chatgpt_response = response_data['choices'][0]['message']['content']
        return chatgpt_response
    else:
        print("Error:", response.status_code, response.text)
        return Non

@app.route('/dashboard')
def dashboard():
    if 'email' not in session:
        return redirect(url_for('signin'))
    
    return render_template('dashboard.html')

if __name__ == '__main__':
    app.run()
